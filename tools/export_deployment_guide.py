from __future__ import annotations

import re
import tempfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "JUNIOR_DEVOPS_DEPLOYMENT_GUIDE.md"
TXT_OUT = ROOT / "docs" / "JUNIOR_DEVOPS_DEPLOYMENT_GUIDE.txt"
DOCX_OUT = ROOT / "docs" / "JUNIOR_DEVOPS_DEPLOYMENT_GUIDE.docx"


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    return text.rstrip()


def paragraph_with_inline_code(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(20, 20, 20)
        else:
            run = paragraph.add_run(part.replace("**", "").replace("__", ""))
            run.font.name = "Arial"
            run.font.size = Pt(11)
    return paragraph


def is_table_row(line: str) -> bool:
    return bool(re.match(r"^\s*\|.*\|\s*$", line))


def is_table_separator(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.match(r"^:?-{3,}:?$", cell.strip()) for cell in cells)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D9DEE8", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin: int = 130) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "left", "bottom", "right"):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(margin))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = tc_pr.first_child_found_in("w:tcW")
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        tc_pr.append(tc_width)
    tc_width.set(qn("w:w"), str(width_dxa))
    tc_width.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_indent(table, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def paragraph_text(cell, text: str):
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.text = ""
    return paragraph


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_width(table, width_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(width_dxa))
    tbl_width.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    set_table_indent(table)


def column_widths(headers: list[str], page_width_dxa: int) -> list[int]:
    weights_by_header = {
        "Машина": 1.15,
        "VM": 0.95,
        "Компонент": 1.45,
        "Компоненты": 2.8,
        "Сервис": 1.45,
        "Что Делает": 2.45,
        "Когда Проверять": 1.75,
        "Сигнал": 1.55,
        "Действие": 2.55,
        "Шаг": 0.8,
        "Команда": 2.2,
        "Порт": 0.8,
        "Откуда Доступ": 1.55,
        "Для Чего": 2.45,
        "CPU": 0.75,
        "RAM": 0.95,
        "Disk": 1.0,
        "Комментарий": 2.25,
    }
    weights = [weights_by_header.get(header, 1.2) for header in headers]
    total = sum(weights)
    widths = [int(page_width_dxa * weight / total) for weight in weights]
    widths[-1] += page_width_dxa - sum(widths)
    return widths


def is_numeric_header(header: str) -> bool:
    lowered = header.lower()
    numeric_tokens = ("cpu", "ram", "disk", "порт", "rows", "bytes", "samples", "%")
    return any(token in lowered for token in numeric_tokens)


def set_run_font(run, size: float, bold: bool = False, color: str = "172033", mono: bool = False) -> None:
    run.bold = bold
    run.font.name = "Courier New" if mono else "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def add_word_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    headers = rows[0]
    body = rows[1:]
    usable_width_dxa = 10560
    widths = column_widths(headers, usable_width_dxa)

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False
    set_table_width(table, usable_width_dxa)
    repeat_table_header(table.rows[0])
    set_cant_split(table.rows[0])

    for column_index, header in enumerate(headers):
        cell = table.rows[0].cells[column_index]
        set_cell_width(cell, widths[column_index])
        set_cell_margins(cell, 140)
        set_cell_border(cell, "1D4ED8", "8")
        set_cell_shading(cell, "1E3A8A")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = paragraph_text(cell, header)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_numeric_header(header) else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(clean_inline_markdown(header))
        set_run_font(run, 8.6, bold=True, color="FFFFFF")

    for row_index, row_values in enumerate(body):
        cells = table.add_row().cells
        set_cant_split(table.rows[-1])
        for column_index, value in enumerate(row_values):
            cell = cells[column_index]
            set_cell_width(cell, widths[column_index])
            set_cell_margins(cell, 130)
            set_cell_border(cell, "D9DEE8", "5")
            if row_index % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            if re.search(r"\b(RUNNING|UP|OK|true|healthy)\b", value, flags=re.I):
                set_cell_shading(cell, "DCFCE7")
            elif re.search(r"\b(DOWN|FAILED|ERROR|false|critical)\b", value, flags=re.I):
                set_cell_shading(cell, "FEE2E2")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = paragraph_text(cell, value)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_numeric_header(headers[column_index]) else WD_ALIGN_PARAGRAPH.LEFT
            parts = re.split(r"(`[^`]+`)", value)
            for part in parts:
                if not part:
                    continue
                if part.startswith("`") and part.endswith("`"):
                    run = paragraph.add_run(part[1:-1])
                    set_run_font(run, 7.7, mono=True, color="0F172A")
                else:
                    run = paragraph.add_run(part.replace("**", "").replace("__", ""))
                    set_run_font(run, 8.0, color="172033")

    document.add_paragraph()


def resolve_image_path(path_text: str) -> Path:
    path = Path(path_text)
    if path.is_absolute():
        return path
    source_relative = (SOURCE.parent / path).resolve()
    if source_relative.exists():
        return source_relative
    return (ROOT / path).resolve()


def add_image(document: Document, alt: str, image_path: Path) -> None:
    if not image_path.exists():
        paragraph_with_inline_code(document, f"[image not found: {image_path}]")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(9.3))
    except Exception:
        with Image.open(image_path) as image:
            normalized = image.convert("RGB")
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as handle:
                normalized.save(handle.name, "PNG")
                normalized_path = handle.name
        run.add_picture(normalized_path, width=Inches(8.2))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    caption_run = caption.add_run(clean_inline_markdown(alt) or image_path.name)
    caption_run.italic = True
    caption_run.font.name = "Arial"
    caption_run.font.size = Pt(8.5)
    caption_run.font.color.rgb = RGBColor(91, 100, 117)


def add_horizontal_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9DEE8")


def markdown_to_txt(markdown: str) -> str:
    lines: list[str] = []
    in_code = False

    for raw in markdown.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            in_code = not in_code
            lines.append("")
            continue

        if re.match(r"^\s*[-_]{3,}\s*$", line):
            if lines and lines[-1] != "":
                lines.append("")
            lines.append("-" * 72)
            lines.append("")
            continue

        if in_code:
            lines.append(line)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            title = clean_inline_markdown(heading.group(2))
            if lines and lines[-1] != "":
                lines.append("")
            lines.append(title.upper() if level == 1 else title)
            lines.append("=" * min(len(title), 80) if level <= 2 else "-" * min(len(title), 80))
            continue

        if re.match(r"^\s*[-*]\s+", line):
            lines.append(re.sub(r"^\s*[-*]\s+", "- ", clean_inline_markdown(line)))
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            lines.append(clean_inline_markdown(line))
            continue

        lines.append(clean_inline_markdown(line))

    text = "\n".join(lines)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip() + "\n"


def markdown_to_docx(markdown: str):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(12)

    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    footer.add_run("Agentic Data Stack deployment guide")

    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[list[str]] = []

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6)
        for index, code_line in enumerate(code_buffer):
            if index:
                paragraph.add_run().add_break(WD_BREAK.LINE)
            run = paragraph.add_run(code_line)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
        code_buffer = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        add_word_table(document, table_buffer)
        table_buffer = []

    for raw in markdown.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
            in_code = not in_code
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if re.match(r"^\s*[-_]{3,}\s*$", line):
            flush_table()
            add_horizontal_rule(document)
            continue

        if not line.strip():
            flush_table()
            continue

        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if image:
            flush_table()
            add_image(document, image.group(1), resolve_image_path(image.group(2)))
            continue

        if is_table_row(line):
            if is_table_separator(line):
                continue
            table_buffer.append(parse_table_row(line))
            continue
        flush_table()

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            title = clean_inline_markdown(heading.group(2))
            level = len(heading.group(1))
            if level == 1:
                document.add_paragraph(title, style="Title")
            else:
                document.add_heading(title, level=min(level, 3))
            continue

        if line.startswith("> "):
            paragraph = paragraph_with_inline_code(document, line[2:])
            for run in paragraph.runs:
                run.italic = True
            continue

        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            paragraph_with_inline_code(document, text, style="List Bullet")
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            paragraph_with_inline_code(document, text, style="List Number")
            continue

        paragraph_with_inline_code(document, line)

    flush_table()
    flush_code()
    return document


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    TXT_OUT.write_text(markdown_to_txt(markdown), encoding="utf-8")
    document = markdown_to_docx(markdown)
    document.save(DOCX_OUT)
    print(TXT_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
